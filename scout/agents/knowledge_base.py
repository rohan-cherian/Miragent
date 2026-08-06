"""
KnowledgeBase — document ingestion and retrieval for the I2R Agent Layer.

Documents are chunked into ~400-word segments and stored in SQLite.
Search uses keyword scoring (TF-style: count matching query terms per chunk).
No external embedding service required — works out of the box.

Upgrade path: swap _score_chunk() for a Weaviate vector query when an
embedding model is available.
"""

from __future__ import annotations

import hashlib
import logging
import re
import string
from typing import Optional

from sqlalchemy import select
from sqlalchemy.orm import Session

from scout.db.models import KBChunk, KBDocument

logger = logging.getLogger(__name__)

STOPWORDS = {
    "the", "a", "an", "is", "are", "was", "were", "be", "been", "being",
    "have", "has", "had", "do", "does", "did", "will", "would", "could",
    "should", "may", "might", "shall", "can", "to", "of", "in", "for",
    "on", "with", "at", "by", "from", "as", "or", "and", "not", "this",
    "that", "it", "its", "we", "our", "you", "your", "they", "their",
}


class KnowledgeBase:
    CHUNK_SIZE = 400      # words per chunk
    CHUNK_OVERLAP = 50    # words of overlap between chunks

    def ingest(
        self,
        content: str,
        filename: str,
        tenant_id: str,
        doc_type: str,
        db: Session,
        created_by: str = "",
    ) -> KBDocument:
        """
        Ingest a text document into the knowledge base.

        Steps:
          1. Hash content for dedup — if same hash exists for tenant, return existing
          2. Chunk content into overlapping 400-word windows
          3. Save KBDocument row
          4. Save KBChunk rows (one per chunk)
          5. Return KBDocument
        """
        content_hash = hashlib.sha256(content.encode("utf-8", errors="replace")).hexdigest()

        # Dedup check
        existing = db.execute(
            select(KBDocument).where(
                KBDocument.tenant_id == tenant_id,
                KBDocument.content_hash == content_hash,
            )
        ).scalar_one_or_none()
        if existing:
            logger.info("Document already ingested (hash match): %s", filename)
            return existing

        chunks = self._chunk_text(content)

        doc = KBDocument(
            tenant_id=tenant_id,
            filename=filename,
            doc_type=doc_type,
            content_hash=content_hash,
            chunk_count=len(chunks),
            created_by=created_by or None,
        )
        db.add(doc)
        db.flush()  # get doc.id before creating chunks

        for i, chunk_text in enumerate(chunks):
            word_count = len(chunk_text.split())
            chunk = KBChunk(
                document_id=doc.id,
                tenant_id=tenant_id,
                chunk_index=i,
                content=chunk_text,
                word_count=word_count,
            )
            db.add(chunk)

        db.commit()
        db.refresh(doc)
        logger.info("Ingested document %s — %d chunks", filename, len(chunks))
        return doc

    def ingest_bytes(
        self,
        raw_bytes: bytes,
        filename: str,
        tenant_id: str,
        doc_type: str,
        db: Session,
        created_by: str = "",
    ) -> KBDocument:
        """Ingest raw file bytes — handles PDF, DOCX, and plain text."""
        content = self._extract_text(raw_bytes, filename)
        return self.ingest(content, filename, tenant_id, doc_type, db, created_by)

    def search(self, query: str, tenant_id: str, db: Session, k: int = 6) -> list[KBChunk]:
        """
        Keyword-scored search over the tenant's KB chunks.

        Steps:
          1. Tokenize query into keywords (lowercase, strip punctuation, remove stopwords)
          2. Load all chunks for tenant from DB
          3. Score each chunk: count of query keyword occurrences in chunk.content
          4. Return top k chunks by score (score > 0 only)
        """
        keywords = self._tokenize(query)
        if not keywords:
            return []

        all_chunks = db.execute(
            select(KBChunk).where(KBChunk.tenant_id == tenant_id)
        ).scalars().all()

        scored: list[tuple[int, KBChunk]] = []
        for chunk in all_chunks:
            score = self._score_chunk(chunk.content, keywords)
            if score > 0:
                scored.append((score, chunk))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [chunk for _, chunk in scored[:k]]

    def list_documents(self, tenant_id: str, db: Session) -> list[KBDocument]:
        """Return all KBDocument rows for tenant, ordered by created_at desc."""
        return db.execute(
            select(KBDocument)
            .where(KBDocument.tenant_id == tenant_id)
            .order_by(KBDocument.created_at.desc())
        ).scalars().all()

    def delete_document(self, doc_id: str, tenant_id: str, db: Session) -> bool:
        """
        Delete a KBDocument and all its KBChunks.

        Returns True if deleted, False if not found.
        """
        doc = db.execute(
            select(KBDocument).where(
                KBDocument.id == doc_id,
                KBDocument.tenant_id == tenant_id,
            )
        ).scalar_one_or_none()

        if not doc:
            return False

        # Chunks are cascade-deleted via the relationship
        db.delete(doc)
        db.commit()
        return True

    def _chunk_text(self, text: str) -> list[str]:
        """
        Split text into overlapping word windows.

        Each chunk is CHUNK_SIZE words; windows shift by (CHUNK_SIZE - CHUNK_OVERLAP).
        """
        words = text.split()
        if not words:
            return []

        step = max(1, self.CHUNK_SIZE - self.CHUNK_OVERLAP)
        chunks: list[str] = []
        i = 0
        while i < len(words):
            chunk_words = words[i : i + self.CHUNK_SIZE]
            chunks.append(" ".join(chunk_words))
            if i + self.CHUNK_SIZE >= len(words):
                break
            i += step
        return chunks

    def _score_chunk(self, chunk_content: str, keywords: list[str]) -> int:
        """Count total occurrences of all keywords in chunk_content (case-insensitive)."""
        content_lower = chunk_content.lower()
        total = 0
        for kw in keywords:
            total += content_lower.count(kw)
        return total

    def _tokenize(self, text: str) -> list[str]:
        """Lowercase, strip punctuation, remove stopwords."""
        text_lower = text.lower()
        # Replace punctuation with spaces
        text_clean = re.sub(r"[^\w\s]", " ", text_lower)
        tokens = text_clean.split()
        return [t for t in tokens if t and t not in STOPWORDS and len(t) > 1]

    def _extract_text(self, raw_bytes: bytes, filename: str) -> str:
        """
        Extract text from raw bytes.

        Handles:
          - .pdf via pypdf (optional — caught gracefully if not installed)
          - .docx via python-docx (optional — caught gracefully if not installed)
          - Plain text (UTF-8, then latin-1 fallback)
        """
        fname_lower = filename.lower()

        if fname_lower.endswith(".pdf"):
            try:
                import io
                import pypdf  # type: ignore[import]

                reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
                pages = []
                for page in reader.pages:
                    pages.append(page.extract_text() or "")
                text = "\n".join(pages)
                if text.strip():
                    return text
                logger.warning("pypdf extracted empty text from %s — falling back to utf-8", filename)
            except ImportError:
                logger.warning("pypdf not installed — treating PDF as plain text")
            except Exception as exc:
                logger.warning("pypdf failed on %s: %s — falling back", filename, exc)

        if fname_lower.endswith(".docx"):
            try:
                import io
                import docx  # type: ignore[import]

                doc = docx.Document(io.BytesIO(raw_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = "\n".join(paragraphs)
                if text.strip():
                    return text
                logger.warning("python-docx extracted empty text from %s — falling back", filename)
            except ImportError:
                logger.warning("python-docx not installed — treating DOCX as plain text")
            except Exception as exc:
                logger.warning("python-docx failed on %s: %s — falling back", filename, exc)

        # Plain text: try UTF-8 first, then latin-1
        try:
            return raw_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return raw_bytes.decode("latin-1")
