"""
DDQTemplateAgent — Sprint 71

Handles DOCX template parsing and filling for investor DDQ forms.

Workflow:
  1. Accept a DOCX file (or plain text) containing an investor DDQ template
  2. Parse question fields from the template
  3. Draft answers using hardcoded high-quality Miragent mock data
  4. Return a DDQTemplateResult including a filled DOCX (base64-encoded)

python-docx is used when available for proper DOCX creation.
Falls back to a base64-encoded plain-text document if unavailable.
"""

from __future__ import annotations

import base64
import io
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Optional
from uuid import uuid4

logger = logging.getLogger(__name__)

# ── Sample DDQ template ────────────────────────────────────────────────────────

SAMPLE_TEMPLATE = """\
DUE DILIGENCE QUESTIONNAIRE
[Company Name]
Submitted by: Acme Capital Partners

1. Describe the company's core product or service offering.

Answer: _______________________________________________

2. What is the company's current ARR and MoM growth rate?

Answer: _______________________________________________

3. Who are the key members of the founding team and what are their backgrounds?

Answer: _______________________________________________

4. What is the company's primary customer acquisition channel?

Answer: _______________________________________________

5. Describe the competitive landscape and how you differentiate.

Answer: _______________________________________________

6. What is the current burn rate and cash runway?

Answer: _______________________________________________

7. How many customers does the company currently serve, and what is the average contract value?

Answer: _______________________________________________

8. What are the key risks the business faces in the next 12 months?

Answer: _______________________________________________

9. Describe the company's go-to-market strategy for the next 12 months.

Answer: _______________________________________________

10. What are the intended uses of the investment proceeds?

Answer: _______________________________________________
"""

# ── Hardcoded high-quality answers for Miragent mock data ─────────────────────

TEMPLATE_ANSWERS: dict[int, str] = {
    1: (
        "Miragent is an AI-powered operational intelligence platform that connects to a company's "
        "existing business systems (CRM, ERP, HRIS), builds a unified knowledge graph, and deploys "
        "specialized AI workers to surface operational risks and opportunities — then autonomously "
        "resolves them through an approval-gated action layer."
    ),
    2: (
        "The company is at $8.42M ARR, growing at 3.3% MoM ($270K net new ARR in April 2026). "
        "Gross retention is 91% and net revenue retention is 107%, indicating healthy expansion revenue."
    ),
    3: (
        "The founding team comprises experienced operators from enterprise SaaS and AI. The CEO brings "
        "12 years of B2B SaaS experience including one prior exit. The CTO leads a team of 18 engineers "
        "with deep expertise in graph databases and agentic AI systems."
    ),
    4: (
        "Primary acquisition is direct enterprise sales (67% of new ARR), supplemented by partnerships "
        "with Big 4 advisory firms and PE/growth equity sponsors who deploy Miragent across portfolio "
        "companies (33% of new ARR). Average sales cycle is 68 days."
    ),
    5: (
        "The market includes point-solution BI tools (Tableau, Looker) and RPA platforms (UiPath). "
        "Miragent differentiates through its read-write architecture: competitors surface insights but "
        "cannot act on them. Our digital twin + autonomous agent layer makes Miragent the only platform "
        "that both diagnoses and remedies operational issues."
    ),
    6: (
        "Monthly cash burn is $318K (down from $342K in March 2026), with $6.84M cash on hand, yielding "
        "21.5 months of runway at current trajectory. The company is burn-efficient relative to ARR, with "
        "a LTV/CAC ratio above 5x."
    ),
    7: (
        "The company serves 127 customers (up from 124 last month), with an average contract value of "
        "$66,300. Customers range from 50-person growth companies to PE-backed portfolio companies with "
        "500+ employees. Top decile customers average $180K ACV."
    ),
    8: (
        "Key risks include: (1) pipeline concentration — top 3 deals represent 31% of weighted pipeline; "
        "(2) hiring velocity — 7 open engineering roles with 18-week average time-to-fill may compress "
        "feature velocity in Q3; (3) AI model dependency — reliance on Anthropic Claude requires "
        "proactive model-provider diversification."
    ),
    9: (
        "The GTM strategy for the next 12 months focuses on three vectors: (1) deepening the PE/sponsor "
        "channel by securing 2-3 additional fund-level partnerships; (2) launching a self-serve tier for "
        "companies under 100 employees; (3) expanding internationally into the UK market through a "
        "London-based enterprise sales hire."
    ),
    10: (
        "Investment proceeds will be allocated as follows: 55% to sales and marketing headcount expansion "
        "(target: double GTM team from 10 to 20); 30% to R&D (AI model improvements, new connector "
        "integrations); 15% to G&A infrastructure and compliance (SOC 2 Type II, international expansion legal)."
    ),
}

# Default answer for questions beyond the hardcoded set
DEFAULT_ANSWER = (
    "Miragent's team is actively reviewing this question and will provide a detailed response "
    "upon request. Please contact your relationship manager for further information."
)


# ── Data classes ───────────────────────────────────────────────────────────────


@dataclass
class DDQTemplateQuestion:
    number: int
    question_text: str
    answer: str
    confidence: float
    confidence_label: str   # "High" / "Review" / "Unknown"


@dataclass
class DDQTemplateResult:
    job_id: str
    tenant_id: str
    source_filename: str
    total_questions: int
    filled_questions: int
    questions: list[DDQTemplateQuestion] = field(default_factory=list)
    filled_docx_b64: str = ""
    can_download: bool = False
    submitted_at: str = ""
    summary: str = ""


# ── Agent ──────────────────────────────────────────────────────────────────────


class DDQTemplateAgent:
    """Parse an investor DDQ template, draft answers, and return a filled DOCX."""

    def process_template(
        self,
        filename: str,
        file_content: bytes,
        tenant_id: str,
        db,
    ) -> DDQTemplateResult:
        """
        Process a DOCX/TXT/PDF file upload.

        Extracts text from the file, parses questions, drafts answers,
        and builds a filled DOCX result.
        """
        text = self._extract_text_from_docx(file_content) if filename.lower().endswith(".docx") else ""
        if not text.strip():
            # Fallback: try raw UTF-8 decode (for .txt, or if docx extraction failed)
            try:
                text = file_content.decode("utf-8")
            except UnicodeDecodeError:
                text = file_content.decode("latin-1")

        result = self.process_text_template(text, tenant_id, db)
        result.source_filename = filename
        return result

    def process_text_template(
        self,
        text: str,
        tenant_id: str,
        db,
    ) -> DDQTemplateResult:
        """
        Process a raw text DDQ template.

        Parses questions, drafts answers using hardcoded Miragent mock data,
        and builds a filled DOCX (or TXT fallback) result.
        """
        job_id = str(uuid4())
        submitted_at = datetime.now(timezone.utc).isoformat()

        raw_questions = self._extract_questions(text)
        if not raw_questions:
            # If no numbered questions found, treat entire text as a single question
            raw_questions = [(1, text.strip())] if text.strip() else []

        questions: list[DDQTemplateQuestion] = []
        for q_num, q_text in raw_questions:
            tq = self._answer_question(q_num, q_text)
            questions.append(tq)

        filled_count = sum(1 for q in questions if q.confidence >= 0.40)
        filled_docx_b64 = self._build_filled_docx_b64(questions)
        can_download = bool(filled_docx_b64)

        summary = (
            f"Parsed {len(questions)} question{'s' if len(questions) != 1 else ''} from template. "
            f"{filled_count} answered with high or review confidence. "
            f"{'DOCX ready for download.' if can_download else 'Plain-text version available.'}"
        )

        return DDQTemplateResult(
            job_id=job_id,
            tenant_id=tenant_id,
            source_filename="sample_ddq_template.txt",
            total_questions=len(questions),
            filled_questions=filled_count,
            questions=questions,
            filled_docx_b64=filled_docx_b64,
            can_download=can_download,
            submitted_at=submitted_at,
            summary=summary,
        )

    def get_sample_template_text(self) -> str:
        """Return the built-in sample DDQ template text."""
        return SAMPLE_TEMPLATE

    # ── Internal helpers ───────────────────────────────────────────────────────

    def _extract_text_from_docx(self, content: bytes) -> str:
        """
        Extract plain text from a DOCX file using python-docx.

        Returns empty string if python-docx is unavailable or extraction fails.
        """
        try:
            import docx  # type: ignore[import]
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            paragraphs.append(cell_text)
            return "\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx not available — falling back to raw bytes decode")
            return ""
        except Exception as exc:
            logger.warning("DOCX text extraction failed: %s", exc)
            return ""

    def _extract_questions(self, text: str) -> list[tuple[int, str]]:
        """
        Parse numbered questions from template text.

        Detects formats:
          - "1. Question text"
          - "1) Question text"
          - "Q1: Question text"
          - "Q1. Question text"

        Also strips trailing "Answer: ___" lines from question text.
        Returns list of (question_number, question_text) tuples.
        """
        numbered_pattern = re.compile(
            r"(?:^|\n)\s*(?:Q(\d+)[:\.]|(\d+)[\.\)])\s+",
            re.MULTILINE,
        )

        matches = list(numbered_pattern.finditer(text))
        if not matches:
            return []

        results: list[tuple[int, str]] = []
        for i, match in enumerate(matches):
            # Extract question number from either group
            q_num_str = match.group(1) or match.group(2)
            q_num = int(q_num_str) if q_num_str else (i + 1)

            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            raw_text = text[start:end].strip()

            # Remove trailing "Answer: ___" lines
            raw_text = re.sub(
                r"\n\s*Answer:\s*[_\s]*$",
                "",
                raw_text,
                flags=re.IGNORECASE | re.MULTILINE,
            ).strip()

            # Also remove standalone "Answer: ___" if it's on the same block
            raw_text = re.sub(
                r"\s*Answer:\s*[_\s]*",
                "",
                raw_text,
                flags=re.IGNORECASE,
            ).strip()

            if raw_text:
                results.append((q_num, raw_text))

        return results

    def _answer_question(self, q_num: int, q_text: str) -> DDQTemplateQuestion:
        """
        Draft an answer for a single question.

        Uses hardcoded TEMPLATE_ANSWERS for questions 1-10.
        Applies a heuristic keyword match for questions outside that range.
        """
        if q_num in TEMPLATE_ANSWERS:
            answer = TEMPLATE_ANSWERS[q_num]
            confidence = 0.92
            confidence_label = "High"
        else:
            # Try keyword matching against our known answers
            q_lower = q_text.lower()
            answer = DEFAULT_ANSWER
            confidence = 0.30
            confidence_label = "Unknown"

            keyword_map = {
                ("product", "service", "offering", "platform"): 1,
                ("arr", "revenue", "growth", "mrr"): 2,
                ("team", "founder", "ceo", "cto", "background"): 3,
                ("acquisition", "channel", "sales", "marketing"): 4,
                ("competitive", "differentiat", "competitor", "landscape"): 5,
                ("burn", "runway", "cash", "spend"): 6,
                ("customer", "contract", "acv", "serve"): 7,
                ("risk", "challenge", "threat"): 8,
                ("go-to-market", "gtm", "strategy", "12 month"): 9,
                ("proceed", "use of fund", "invest", "allocat"): 10,
            }
            for kws, num in keyword_map.items():
                if any(kw in q_lower for kw in kws):
                    answer = TEMPLATE_ANSWERS[num]
                    confidence = 0.72
                    confidence_label = "Review"
                    break

        return DDQTemplateQuestion(
            number=q_num,
            question_text=q_text,
            answer=answer,
            confidence=confidence,
            confidence_label=confidence_label,
        )

    def _build_filled_docx_b64(self, questions: list[DDQTemplateQuestion]) -> str:
        """
        Build a filled DOCX document and return it as a base64-encoded string.

        Uses python-docx if available; falls back to a plain-text .txt in b64.
        """
        date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        header = f"COMPLETED DDQ — Generated by Miragent — {date_str}"

        try:
            from docx import Document  # type: ignore[import]
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]

            doc = Document()

            # Header
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(header)
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

            doc.add_paragraph()  # Spacer

            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run("DUE DILIGENCE QUESTIONNAIRE — Miragent, Inc.")
            sub_run.bold = True
            sub_run.font.size = Pt(12)

            doc.add_paragraph()  # Spacer

            for tq in questions:
                # Question label (bold)
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Q{tq.number}. {tq.question_text}")
                q_run.bold = True
                q_run.font.size = Pt(11)

                # Answer text (normal)
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(tq.answer)
                a_run.font.size = Pt(10)

                # Confidence footnote (small, gray)
                conf_para = doc.add_paragraph()
                conf_run = conf_para.add_run(
                    f"[Confidence: {tq.confidence_label} — {tq.confidence * 100:.0f}%]"
                )
                conf_run.font.size = Pt(8)
                conf_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                doc.add_paragraph()  # Spacer between Q&A pairs

            buf = io.BytesIO()
            doc.save(buf)
            return base64.b64encode(buf.getvalue()).decode("ascii")

        except ImportError:
            logger.info("python-docx not available — generating plain-text fallback")
        except Exception as exc:
            logger.warning("DOCX generation failed: %s — falling back to plain text", exc)

        # Plain-text fallback
        lines = [header, "", "DUE DILIGENCE QUESTIONNAIRE — Miragent, Inc.", ""]
        for tq in questions:
            lines.append(f"Q{tq.number}. {tq.question_text}")
            lines.append("")
            lines.append(f"Answer: {tq.answer}")
            lines.append(f"[Confidence: {tq.confidence_label} — {tq.confidence * 100:.0f}%]")
            lines.append("")
            lines.append("-" * 80)
            lines.append("")

        plain_text = "\n".join(lines)
        return base64.b64encode(plain_text.encode("utf-8")).decode("ascii")
